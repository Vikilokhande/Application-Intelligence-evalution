"""
extraction/providers.py
=======================

Real provider implementations for document parsing, OCR, LLM, and embeddings.

Production dependency graph:
  DocumentParser  → RealDocumentParser
  OCRProvider     → TesseractOCRProvider
  LLMProvider     → GroqLLMProvider
  EmbeddingProvider → SentenceTransformerEmbeddingProvider (or LocalFallbackEmbeddingProvider)

Test-only mock providers live under backend/tests/fixtures, not in application code.
"""

from __future__ import annotations

import csv
import email.utils
import io
import json
import logging
import re
import shutil
import time
import urllib.error
import urllib.request
from abc import ABC, abstractmethod
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)


def _parse_retry_after(value: Any) -> float | None:
    """Parse Retry-After header seconds or HTTP-date into a non-negative delay."""
    if value is None:
        return None
    raw = str(value).strip()
    if not raw:
        return None
    try:
        return max(0.0, float(raw))
    except (TypeError, ValueError):
        pass
    try:
        retry_at = email.utils.parsedate_to_datetime(raw)
        if retry_at.tzinfo is None:
            retry_at = retry_at.replace(tzinfo=timezone.utc)
        return max(0.0, (retry_at - datetime.now(timezone.utc)).total_seconds())
    except (TypeError, ValueError, IndexError, OverflowError):
        return None


def _response_retry_after(exc: Exception) -> float | None:
    # 1. Check HTTP headers on response or exception
    response = getattr(exc, "response", None)
    headers = getattr(response, "headers", None)
    if headers is None:
        headers = getattr(exc, "headers", None)
    if headers is not None:
        getter = getattr(headers, "get", None)
        if callable(getter):
            val = _parse_retry_after(getter("Retry-After") or getter("retry-after"))
            if val is not None:
                return val

    # 2. Check OpenAI API body / exception JSON metadata (common for OpenRouter 429)
    body = getattr(exc, "body", None)
    if isinstance(body, dict):
        err = body.get("error", {})
        if isinstance(err, dict):
            meta = err.get("metadata", {})
            if isinstance(meta, dict):
                sec = meta.get("retry_after_seconds") or meta.get("retry_after")
                if sec is not None:
                    try:
                        return max(0.0, float(sec))
                    except (ValueError, TypeError):
                        pass
                hdrs = meta.get("headers", {})
                if isinstance(hdrs, dict):
                    h_val = hdrs.get("Retry-After") or hdrs.get("retry-after")
                    if h_val is not None:
                        val = _parse_retry_after(h_val)
                        if val is not None:
                            return val
    return None


def _infer_status_code(message: str) -> int | None:
    match = re.search(r"\b([45][0-9]{2})\b", message)
    return int(match.group(1)) if match else None


def detect_tesseract_cmd(configured_path: str = "") -> str:
    """Return a usable Tesseract executable path, checking config, PATH, then common Windows locations."""
    candidates = []
    if configured_path:
        candidates.append(configured_path)
    path_hit = shutil.which("tesseract")
    if path_hit:
        candidates.append(path_hit)
    candidates.extend(
        [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        ]
    )
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return candidate
    return configured_path

# ---------------------------------------------------------------------------
# Abstract Interfaces
# ---------------------------------------------------------------------------


class DocumentParser(ABC):
    @abstractmethod
    def parse(self, path: Path, mime_type: str) -> dict[str, Any]:
        """Parse a document and return normalised text + metadata."""
        raise NotImplementedError


class OCRProvider(ABC):
    @abstractmethod
    def extract_text(self, path: Path, language: str = "eng") -> dict[str, Any]:
        """Run OCR on path; return {text, confidence, page_count, provider, status}."""
        raise NotImplementedError


class LLMProvider(ABC):
    @abstractmethod
    def generate(self, prompt: str, system: str = "") -> str:
        raise NotImplementedError

    @abstractmethod
    def summarize(self, text: str) -> str:
        raise NotImplementedError

    @abstractmethod
    def extract_structured(self, text: str, schema_name: str) -> dict[str, Any]:
        raise NotImplementedError

    @abstractmethod
    def classify_document(self, text: str, filename: str, declared_type: str) -> dict[str, Any]:
        raise NotImplementedError


class EmbeddingProvider(ABC):
    @abstractmethod
    def embed(self, texts: list[str]) -> list[list[float]]:
        raise NotImplementedError


# ---------------------------------------------------------------------------
# Real Document Parser
# ---------------------------------------------------------------------------


class RealDocumentParser(DocumentParser):
    """
    Parse documents using format-appropriate real libraries.

    Supported:
      .pdf   → pypdf
      .docx  → python-docx
      .xlsx  → openpyxl
      .csv   → stdlib csv
      .json  → stdlib json
      .txt   → pathlib read_text
      images → (passed to OCR separately; returns empty text here)
    """

    def parse(self, path: Path, mime_type: str) -> dict[str, Any]:
        suffix = path.suffix.lower()
        try:
            if suffix == ".pdf":
                return self._parse_pdf(path)
            if suffix == ".docx":
                return self._parse_docx(path)
            if suffix in (".xlsx", ".xls"):
                return self._parse_xlsx(path)
            if suffix == ".csv":
                return self._parse_csv(path)
            if suffix == ".json":
                return self._parse_json(path)
            if suffix in (".txt", ".md"):
                return self._parse_text(path)
            if suffix in (".jpg", ".jpeg", ".png", ".tiff", ".tif"):
                # Image files → text extracted by OCR provider; return stub here
                return {"text": "", "pages": [], "tables": [], "metadata": {"filename": path.name, "mime_type": mime_type, "requires_ocr": True}, "parser": "none", "parser_version": "0"}
            # Fallback: try reading as text
            return self._parse_text(path)
        except Exception as exc:
            logger.warning("Document parsing failed for %s: %s", path.name, exc)
            return {"text": "", "pages": [], "tables": [], "metadata": {"filename": path.name, "error": str(exc)}, "parser": "failed", "parser_version": "0"}

    # ---- PDF ----------------------------------------------------------------

    def _parse_pdf(self, path: Path) -> dict[str, Any]:
        try:
            import pypdf  # type: ignore
        except ImportError as exc:
            raise ImportError("pypdf is required for PDF parsing. Install with: pip install pypdf") from exc

        try:
            reader = pypdf.PdfReader(str(path))
            pages: list[dict[str, Any]] = []
            all_text_parts: list[str] = []

            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                pages.append({"page": page_num, "text": page_text, "char_count": len(page_text)})
                if page_text.strip():
                    all_text_parts.append(f"[Page {page_num}]\n{page_text}")

            full_text = "\n\n".join(all_text_parts)
            return {
                "text": full_text,
                "pages": pages,
                "tables": [],
                "metadata": {
                    "filename": path.name,
                    "page_count": len(reader.pages),
                    "size_bytes": path.stat().st_size,
                },
                "parser": "pypdf",
                "parser_version": getattr(pypdf, "__version__", "unknown"),
            }
        except Exception:
            # Fallback if text-based file was saved with .pdf extension
            try:
                raw_text = path.read_text(encoding="utf-8", errors="ignore").strip()
                if len(raw_text) > 10 and not raw_text.startswith("%PDF"):
                    return {
                        "text": raw_text,
                        "pages": [{"page": 1, "text": raw_text, "char_count": len(raw_text)}],
                        "tables": [],
                        "metadata": {"filename": path.name, "size_bytes": path.stat().st_size},
                        "parser": "text-fallback",
                        "parser_version": "1.0",
                    }
            except Exception:
                pass
            raise

    # ---- DOCX ---------------------------------------------------------------

    def _parse_docx(self, path: Path) -> dict[str, Any]:
        try:
            import docx  # type: ignore
        except ImportError as exc:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx") from exc

        doc = docx.Document(str(path))
        paragraphs: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]

        tables: list[dict[str, Any]] = []
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            tables.append({"table_index": table_idx, "rows": rows})

        table_texts: list[str] = []
        for t in tables:
            for row in t["rows"]:
                table_texts.append(" | ".join(row))

        full_text = "\n".join(paragraphs)
        if table_texts:
            full_text += "\n\n[Tables]\n" + "\n".join(table_texts)

        return {
            "text": full_text,
            "pages": [{"page": 1, "text": full_text, "char_count": len(full_text)}],
            "tables": tables,
            "metadata": {"filename": path.name, "size_bytes": path.stat().st_size},
            "parser": "python-docx",
            "parser_version": getattr(docx, "__version__", "unknown"),
        }

    # ---- XLSX ---------------------------------------------------------------

    def _parse_xlsx(self, path: Path) -> dict[str, Any]:
        try:
            import openpyxl  # type: ignore
        except ImportError as exc:
            raise ImportError("openpyxl is required for XLSX parsing. Install with: pip install openpyxl") from exc

        wb = openpyxl.load_workbook(str(path), data_only=True)
        all_text_parts: list[str] = []
        tables: list[dict[str, Any]] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows: list[list[str]] = []
            for row in ws.iter_rows():
                cells = [str(cell.value) if cell.value is not None else "" for cell in row]
                if any(c.strip() for c in cells):
                    rows.append(cells)

            tables.append({"sheet": sheet_name, "rows": rows})
            sheet_text = f"[Sheet: {sheet_name}]\n"
            sheet_text += "\n".join(" | ".join(row) for row in rows)
            all_text_parts.append(sheet_text)

        full_text = "\n\n".join(all_text_parts)
        return {
            "text": full_text,
            "pages": [{"page": 1, "text": full_text, "char_count": len(full_text)}],
            "tables": tables,
            "metadata": {"filename": path.name, "sheet_count": len(wb.sheetnames), "size_bytes": path.stat().st_size},
            "parser": "openpyxl",
            "parser_version": getattr(openpyxl, "__version__", "unknown"),
        }

    # ---- CSV ----------------------------------------------------------------

    def _parse_csv(self, path: Path) -> dict[str, Any]:
        content = path.read_text(encoding="utf-8", errors="ignore")
        reader = csv.reader(io.StringIO(content))
        rows = list(reader)
        text_lines = [" | ".join(row) for row in rows if any(c.strip() for c in row)]
        full_text = "\n".join(text_lines)
        return {
            "text": full_text,
            "pages": [{"page": 1, "text": full_text, "char_count": len(full_text)}],
            "tables": [{"rows": rows}],
            "metadata": {"filename": path.name, "row_count": len(rows), "size_bytes": path.stat().st_size},
            "parser": "stdlib-csv",
            "parser_version": "3.x",
        }

    # ---- JSON ---------------------------------------------------------------

    def _parse_json(self, path: Path) -> dict[str, Any]:
        content = path.read_text(encoding="utf-8", errors="ignore")
        try:
            data = json.loads(content)
            text = json.dumps(data, indent=2, ensure_ascii=False)
        except json.JSONDecodeError:
            text = content
        return {
            "text": text,
            "pages": [{"page": 1, "text": text, "char_count": len(text)}],
            "tables": [],
            "metadata": {"filename": path.name, "size_bytes": path.stat().st_size},
            "parser": "stdlib-json",
            "parser_version": "3.x",
        }

    # ---- Plain text ---------------------------------------------------------

    def _parse_text(self, path: Path) -> dict[str, Any]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return {
            "text": text,
            "pages": [{"page": 1, "text": text, "char_count": len(text)}],
            "tables": [],
            "metadata": {"filename": path.name, "size_bytes": path.stat().st_size},
            "parser": "pathlib-text",
            "parser_version": "3.x",
        }


# ---------------------------------------------------------------------------
# Real OCR Provider — Tesseract
# ---------------------------------------------------------------------------


class TesseractOCRProvider(OCRProvider):
    """
    Production OCR using pytesseract + Tesseract binary.

    Raises OCRProviderError if Tesseract is not installed or fails.
    Never returns fake text.
    """

    def __init__(self, tesseract_cmd: str = "", language: str = "eng", timeout: int = 60) -> None:
        self.language = language
        self.timeout = timeout
        self.tesseract_cmd = detect_tesseract_cmd(tesseract_cmd)
        if self.tesseract_cmd:
            try:
                import pytesseract  # type: ignore
                pytesseract.pytesseract.tesseract_cmd = self.tesseract_cmd
            except ImportError:
                pass  # Will be caught in extract_text

    def _check_available(self) -> None:
        try:
            import pytesseract  # type: ignore
            if self.tesseract_cmd:
                pytesseract.pytesseract.tesseract_cmd = self.tesseract_cmd
            pytesseract.get_tesseract_version()
        except ImportError as exc:
            from app.core.exceptions import OCRProviderError
            raise OCRProviderError(
                "pytesseract is not installed. Install with: pip install pytesseract. "
                "Also ensure the Tesseract binary is installed on the system.",
                provider="tesseract",
            ) from exc
        except Exception as exc:
            from app.core.exceptions import OCRProviderError
            raise OCRProviderError(
                f"Tesseract binary not found or not functional: {exc}. "
                "Install Tesseract OCR on your system. "
                "On Windows, set TESSERACT_CMD to the path of tesseract.exe.",
                provider="tesseract",
            ) from exc

    def extract_text(self, path: Path, language: str = "") -> dict[str, Any]:
        from app.core.exceptions import OCRProviderError

        lang = language or self.language
        self._check_available()

        try:
            import pytesseract  # type: ignore
            from PIL import Image  # type: ignore
        except ImportError as exc:
            raise OCRProviderError(
                "Pillow is required for OCR image processing. Install with: pip install Pillow",
                provider="tesseract",
            ) from exc

        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return self._ocr_pdf(path, lang, pytesseract, Image)
        else:
            return self._ocr_image(path, lang, pytesseract, Image)

    def _ocr_image(self, path: Path, lang: str, pytesseract: Any, Image: Any) -> dict[str, Any]:
        from app.core.exceptions import OCRProviderError
        try:
            img = Image.open(str(path))
            data = pytesseract.image_to_data(img, lang=lang, output_type=pytesseract.Output.DICT)
            text = pytesseract.image_to_string(img, lang=lang)
            confidences = [int(c) for c in data["conf"] if str(c).strip() != "-1" and str(c).strip() != ""]
            avg_confidence = round(sum(confidences) / len(confidences) / 100.0, 3) if confidences else 0.0
            return {
                "text": text,
                "pages": [{"page": 1, "text": text}],
                "confidence": avg_confidence,
                "page_count": 1,
                "provider": "tesseract",
                "language": lang,
                "status": "SUCCESS",
            }
        except Exception as exc:
            raise OCRProviderError(
                f"Tesseract OCR failed on image {path.name}: {exc}",
                provider="tesseract",
            ) from exc

    def _ocr_pdf(self, path: Path, lang: str, pytesseract: Any, Image: Any) -> dict[str, Any]:
        """
        Render PDF pages using PyMuPDF (fitz) at 300 DPI and run Tesseract on each page.
        No Poppler / pdf2image dependency required.
        """
        from app.core.exceptions import OCRProviderError

        try:
            import fitz  # PyMuPDF  # type: ignore
        except ImportError as exc:
            raise OCRProviderError(
                "PyMuPDF is required for PDF OCR without Poppler. "
                "Install with: pip install PyMuPDF",
                provider="tesseract",
            ) from exc

        try:
            pdf_doc = fitz.open(str(path))
        except Exception as exc:
            try:
                raw_text = path.read_text(encoding="utf-8", errors="ignore").strip()
                if len(raw_text) > 10:
                    return {
                        "text": raw_text,
                        "pages": [{"page": 1, "text": raw_text, "confidence": 1.0}],
                        "confidence": 1.0,
                        "page_count": 1,
                        "provider": "text-fallback",
                        "renderer": "text",
                        "language": lang,
                        "status": "SUCCESS",
                    }
            except Exception:
                pass
            raise OCRProviderError(
                f"PyMuPDF could not open PDF '{path.name}': {exc}",
                provider="tesseract",
            ) from exc

        pages: list[dict[str, Any]] = []
        all_text: list[str] = []
        all_confidences: list[float] = []

        # 300 DPI rendering matrix (scale = 300/72 ≈ 4.167)
        zoom = 300 / 72
        mat = fitz.Matrix(zoom, zoom)

        for page_num in range(len(pdf_doc)):
            fitz_page = pdf_doc.load_page(page_num)
            try:
                # Render to pixmap → convert to PIL Image
                pixmap = fitz_page.get_pixmap(matrix=mat, alpha=False)
                img = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)

                data = pytesseract.image_to_data(img, lang=lang, output_type=pytesseract.Output.DICT)
                page_text = pytesseract.image_to_string(img, lang=lang)
                confidences = [int(c) for c in data["conf"] if str(c).strip() not in ("-1", "")]
                page_conf = round(sum(confidences) / len(confidences) / 100.0, 3) if confidences else 0.0
                all_confidences.extend([c / 100.0 for c in confidences])
                pages.append({"page": page_num + 1, "text": page_text, "confidence": page_conf})
                if page_text.strip():
                    all_text.append(f"[Page {page_num + 1}]\n{page_text}")
            except Exception as exc:
                logger.warning("OCR failed on page %d of %s: %s", page_num + 1, path.name, exc)
                pages.append({"page": page_num + 1, "text": "", "error": str(exc)})

        pdf_doc.close()
        avg_confidence = round(sum(all_confidences) / len(all_confidences), 3) if all_confidences else 0.0
        full_text = "\n\n".join(all_text)
        return {
            "text": full_text,
            "pages": pages,
            "confidence": avg_confidence,
            "page_count": len(pdf_doc) if hasattr(pdf_doc, "__len__") else len(pages),
            "provider": "tesseract",
            "renderer": "pymupdf",
            "language": lang,
            "status": "SUCCESS",
        }


# ---------------------------------------------------------------------------
# Real LLM Provider — Groq (OpenAI-compatible)
# ---------------------------------------------------------------------------


class GroqLLMProvider(LLMProvider):
    """
    Production LLM using the Groq OpenAI-compatible API.

    - All calls are timeout-guarded.
    - API key is NEVER logged.
    - Structured JSON extraction uses JSON mode.
    - Raises LLMProviderError on failure; never returns fake data.
    """

    EXTRACTION_PROMPT = """You are an AI assistant that extracts structured information from government grant/scheme application documents.

Extract the following fields from the document text below.
Return ONLY a valid JSON object. Do not add any explanation, markdown, or text outside the JSON.

Rules:
1. Use ONLY information explicitly present in the text. NEVER invent or guess values.
2. If a field is not present, set its value to null.
3. Preserve original units (lakh, crore, months, etc.) in the source field.
4. Provide a confidence score (0.0 to 1.0) for each extracted field.
5. For financial values, convert to base currency units (rupees if Indian currency).

Fields to extract:
- applicant_name: string | null
- organization_type: string | null  (e.g. Municipality, Registered NGO, Academic Institution, Government Agency)
- project_title: string | null
- project_category: string | null
- project_cost: number | null  (in rupees)
- duration_months: integer | null
- certificate_number: string | null
- environmental_benefit: string | null

Return format (for each field):
{{
  "applicant_name": {{"value": "...", "confidence": 0.95, "source": "text excerpt", "reason": "explicit match"}},
  ...
}}

Document text:
---
{text}
---"""

    CLASSIFICATION_PROMPT = """Classify this government application document into one of these types:
APPLICATION_FORM, PROPOSAL, BUDGET, CERTIFICATE, TIMELINE, TECHNICAL_REPORT, FINANCIAL_REPORT, SUPPORTING_DOCUMENT, OTHER

Filename: {filename}
Declared type: {declared_type}

Text excerpt (first 1000 chars):
---
{text_excerpt}
---

Return ONLY a JSON object:
{{"document_type": "...", "confidence": 0.0, "reason": "...", "signals": []}}"""

    SUMMARIZE_PROMPT = """Summarize this government application document in 2-3 sentences.
Focus on: applicant, project objective, cost, and duration.
If the text is too short or unclear, say so honestly.

Document text:
---
{text}
---

Return only the summary text, no markdown."""

    def __init__(
        self,
        api_key: str,
        model: str,
        base_url: str,
        temperature: float = 0.0,
        timeout: int = 60,
        max_tokens: int = 4096,
        max_retries: int = 2,
        retry_backoff_seconds: float = 3.0,
    ) -> None:
        if not api_key:
            from app.core.exceptions import LLMProviderError
            raise LLMProviderError("API key is required. Set OPENROUTER_API_KEY in .env.", provider="openrouter")
        self._api_key = api_key  # stored privately; never logged
        self.model = model
        self.base_url = base_url.rstrip("/")
        self.temperature = temperature
        self.timeout = timeout
        self.max_tokens = max_tokens
        self.max_retries = max(0, int(max_retries))
        self.retry_backoff_seconds = max(0.1, float(retry_backoff_seconds))

    def _call_api_legacy(
        self,
        messages: list[dict[str, str]],
        json_mode: bool = False,
        correlation_id: str = "",
        backoff_base: float = 3.0,
    ) -> str:
        """
        Call the LLM API with bounded exponential backoff on 429 rate limits.

        On 429: reads Retry-After header, waits, retries up to self.max_retries times.
        On 401/403: breaks immediately (auth error, no point retrying).
        On other errors: exponential backoff and retry.
        After max_retries exhausted: raises LLMProviderError.
        """
        from app.core.exceptions import LLMProviderError

        last_error: Exception | None = None
        for attempt in range(1, self.max_retries + 1):
            try:
                logger.info(
                    "[LLM] model=%s attempt=%d/%d corr=%s",
                    self.model, attempt, self.max_retries, correlation_id or "-",
                )
                content = self._call_via_openai_sdk(messages, json_mode)
                logger.info(
                    "[LLM] model=%s attempt=%d status=SUCCESS corr=%s",
                    self.model, attempt, correlation_id or "-",
                )
                return content

            except LLMProviderError as exc:
                last_error = exc
                msg = exc.message

                # --- 401 / 403: auth errors — do NOT retry ---
                if "401" in msg or "403" in msg:
                    logger.error(
                        "[LLM] model=%s attempt=%d status=AUTH_ERROR — stopping retries",
                        self.model, attempt,
                    )
                    break

                # --- 429: rate limit — read Retry-After if available ---
                is_rate_limit = "429" in msg or "rate limit" in msg.lower() or "rate_limit" in msg.lower()
                if is_rate_limit:
                    # Try to extract Retry-After from the exception chain
                    retry_after_secs: float | None = None
                    cause = getattr(exc, "__cause__", None)
                    response_obj = getattr(cause, "response", None)
                    if response_obj is not None:
                        header_val = getattr(response_obj, "headers", {}).get("retry-after")
                        if header_val:
                            try:
                                retry_after_secs = float(header_val)
                            except (ValueError, TypeError):
                                pass

                    delay = min(30.0, retry_after_secs if retry_after_secs else backoff_base * (2 ** (attempt - 1)))
                    logger.warning(
                        "[LLM] model=%s attempt=%d status=RATE_LIMITED retry_after=%.1fs",
                        self.model, attempt, delay,
                    )
                    if attempt < self.max_retries:
                        time.sleep(delay)
                    continue

                # --- Other errors: exponential backoff ---
                delay = min(30.0, backoff_base * (2 ** (attempt - 1)))
                logger.warning(
                    "[LLM] model=%s attempt=%d status=ERROR delay=%.1fs error=%.150s",
                    self.model, attempt, delay, msg,
                )
                if attempt < self.max_retries:
                    time.sleep(delay)

            except Exception as exc:
                last_error = exc
                delay = min(30.0, backoff_base * (2 ** (attempt - 1)))
                logger.warning(
                    "[LLM] model=%s attempt=%d status=UNEXPECTED_ERROR delay=%.1fs type=%s",
                    self.model, attempt, delay, type(exc).__name__,
                )
                if attempt < self.max_retries:
                    time.sleep(delay)

        raise LLMProviderError(
            f"LLM model '{self.model}' failed after {self.max_retries} attempt(s): {last_error}",
            provider="openrouter",
        ) from last_error

    def _call_api(
        self,
        messages: list[dict[str, str]],
        json_mode: bool = False,
        correlation_id: str = "",
        backoff_base: float | None = None,
        log_scope: str = "LLM",
    ) -> str:
        """
        Call the LLM API with bounded retry handling.

        max_retries means retries after the initial attempt, so the default of 2
        gives at most 3 attempts per model.
        """
        from app.core.exceptions import LLMProviderError

        last_error: Exception | None = None
        last_status_code: int | None = None
        last_retry_after: float | None = None
        scope = (log_scope or "LLM").strip()
        base_delay = max(0.1, float(backoff_base if backoff_base is not None else self.retry_backoff_seconds))
        max_attempts = self.max_retries + 1

        for attempt in range(1, max_attempts + 1):
            try:
                logger.info(
                    "[%s] model=%s attempt=%d status=STARTED corr=%s",
                    scope,
                    self.model,
                    attempt,
                    correlation_id or "-",
                )
                content = self._call_via_openai_sdk(messages, json_mode)
                logger.info(
                    "[%s] model=%s attempt=%d status=SUCCESS",
                    scope,
                    self.model,
                    attempt,
                )
                return content
            except LLMProviderError as exc:
                last_error = exc
                msg = exc.message
                last_status_code = exc.status_code or _infer_status_code(msg)
                last_retry_after = exc.retry_after
                has_retry = attempt < max_attempts

                if last_status_code in (400, 401, 403, 404) or any(code in msg for code in ("400", "401", "403", "404")):
                    logger.error(
                        "[%s] model=%s attempt=%d status=FAILED status_code=%s",
                        scope,
                        self.model,
                        attempt,
                        last_status_code or "client_error",
                    )
                    break

                is_rate_limit = (
                    last_status_code == 429
                    or "429" in msg
                    or "rate limit" in msg.lower()
                    or "rate_limit" in msg.lower()
                )
                if is_rate_limit:
                    delay = min(
                        30.0,
                        last_retry_after
                        if last_retry_after is not None
                        else base_delay * (2 ** (attempt - 1)),
                    )
                    logger.warning(
                        "[%s] model=%s attempt=%d status=RATE_LIMITED retry_after=%.1f",
                        scope,
                        self.model,
                        attempt,
                        delay,
                    )
                    if has_retry:
                        time.sleep(delay)
                        continue
                    logger.warning(
                        "[%s] model=%s attempt=%d status=FAILED status_code=429",
                        scope,
                        self.model,
                        attempt,
                    )
                    break

                delay = min(30.0, base_delay * (2 ** (attempt - 1)))
                logger.warning(
                    "[%s] model=%s attempt=%d status=FAILED status_code=%s retry_delay=%.1f error=%.150s",
                    scope,
                    self.model,
                    attempt,
                    last_status_code or "unknown",
                    delay if has_retry else 0.0,
                    msg,
                )
                if has_retry:
                    time.sleep(delay)
                    continue
                break
            except Exception as exc:
                last_error = exc
                has_retry = attempt < max_attempts
                delay = min(30.0, base_delay * (2 ** (attempt - 1)))
                logger.warning(
                    "[%s] model=%s attempt=%d status=FAILED status_code=unknown retry_delay=%.1f error_type=%s",
                    scope,
                    self.model,
                    attempt,
                    delay if has_retry else 0.0,
                    type(exc).__name__,
                )
                if has_retry:
                    time.sleep(delay)
                    continue
                break

        raise LLMProviderError(
            f"LLM model '{self.model}' failed after {max_attempts} attempt(s): {last_error}",
            provider="openrouter",
            status_code=last_status_code,
            retry_after=last_retry_after,
            model=self.model,
        ) from last_error

    def _call_via_openai_sdk(self, messages: list[dict[str, str]], json_mode: bool = False) -> str:
        """Use the openai Python SDK which properly handles headers and avoids Cloudflare blocks."""
        from app.core.exceptions import LLMProviderError
        try:
            from openai import OpenAI, APIError, AuthenticationError, RateLimitError  # type: ignore
        except ImportError:
            # Fall back to urllib if openai not installed
            return self._call_via_urllib(messages, json_mode)

        try:
            # OpenRouter requires HTTP-Referer and X-Title headers for free-tier models
            extra_headers: dict[str, str] = {}
            if "openrouter.ai" in self.base_url:
                extra_headers["HTTP-Referer"] = "https://localhost:5173"
                extra_headers["X-Title"] = "Application Intelligence Platform"

            client = OpenAI(
                api_key=self._api_key,
                base_url=self.base_url,
                timeout=self.timeout,
                max_retries=0,  # We handle retries ourselves
                default_headers=extra_headers if extra_headers else None,
            )
            kwargs: dict[str, Any] = {
                "model": self.model,
                "messages": messages,
                "temperature": self.temperature,
                "max_tokens": self.max_tokens,
            }
            # Note: response_format=json_object is NOT set here because not all
            # models support it. JSON fidelity is enforced via prompt instructions,
            # and robust extraction is done in callers.
            response = client.chat.completions.create(**kwargs)
            return response.choices[0].message.content or ""
        except AuthenticationError as exc:
            raise LLMProviderError(
                f"Authentication failed (401): {exc}",
                provider="openrouter",
                status_code=getattr(exc, "status_code", 401) or 401,
                retry_after=_response_retry_after(exc),
                model=self.model,
            ) from exc
        except RateLimitError as exc:
            raise LLMProviderError(
                f"Rate limit exceeded (429): {exc}",
                provider="openrouter",
                status_code=getattr(exc, "status_code", 429) or 429,
                retry_after=_response_retry_after(exc),
                model=self.model,
            ) from exc
        except APIError as exc:
            raise LLMProviderError(
                f"API error ({exc.status_code}): {exc.message}",
                provider="openrouter",
                status_code=getattr(exc, "status_code", None),
                retry_after=_response_retry_after(exc),
                model=self.model,
            ) from exc
        except Exception as exc:
            raise LLMProviderError(
                f"OpenAI SDK error: {exc}",
                provider="openrouter",
                model=self.model,
            ) from exc

    def _call_via_urllib(self, messages: list[dict[str, str]], json_mode: bool = False) -> str:
        """Fallback: raw urllib request (may be blocked by some CDN configurations)."""
        from app.core.exceptions import LLMProviderError
        payload: dict[str, Any] = {
            "model": self.model,
            "messages": messages,
            "temperature": self.temperature,
            "max_tokens": self.max_tokens,
        }
        if json_mode:
            payload["response_format"] = {"type": "json_object"}
        body = json.dumps(payload).encode("utf-8")
        url = f"{self.base_url}/chat/completions"
        try:
            req = urllib.request.Request(
                url, data=body,
                headers={
                    "Authorization": f"Bearer {self._api_key}",
                    "Content-Type": "application/json",
                    "User-Agent": "ApplicationIntelligencePlatform/1.0",
                    # OpenRouter required headers
                    **({"HTTP-Referer": "https://localhost:5173", "X-Title": "Application Intelligence Platform"}
                       if "openrouter.ai" in self.base_url else {}),
                },
                method="POST",
            )
            with urllib.request.urlopen(req, timeout=self.timeout) as resp:
                result = json.loads(resp.read())
            return result["choices"][0]["message"]["content"]
        except urllib.error.HTTPError as exc:
            err_body = exc.read().decode("utf-8", errors="ignore")
            raise LLMProviderError(
                f"HTTP {exc.code}: {err_body[:300]}",
                provider="openrouter",
                status_code=exc.code,
                retry_after=_parse_retry_after(exc.headers.get("Retry-After") if exc.headers else None),
                model=self.model,
            ) from exc
        except Exception as exc:
            raise LLMProviderError(
                f"urllib request failed: {exc}",
                provider="openrouter",
                model=self.model,
            ) from exc

    @staticmethod
    def _extract_json_from_text(text: str) -> dict[str, Any]:
        """
        Robustly extract JSON object from LLM response.
        Handles: plain JSON, markdown code fences, leading text, trailing garbage.
        """
        import re
        # Strip markdown code fences
        text = re.sub(r"```(?:json)?\s*", "", text, flags=re.IGNORECASE).strip()
        text = re.sub(r"```\s*$", "", text).strip()
        # Try direct parse first
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass
        # Find the first {...} block (handles leading prose)
        match = re.search(r"\{[\s\S]*\}", text)
        if match:
            try:
                return json.loads(match.group())
            except json.JSONDecodeError:
                # Try to fix truncated JSON by closing open brackets
                fragment = match.group()
                for suffix in ["}}", "}", '"}', '"}}']:
                    try:
                        return json.loads(fragment + suffix)
                    except json.JSONDecodeError:
                        continue
        raise json.JSONDecodeError(f"No valid JSON object found in LLM response", text, 0)

    def generate(self, prompt: str, system: str = "") -> str:
        messages: list[dict[str, str]] = []
        if system:
            messages.append({"role": "system", "content": system})
        messages.append({"role": "user", "content": prompt})
        return self._call_api(messages)

    def summarize(self, text: str) -> str:
        if not text.strip():
            return "No readable text was extracted from this document."
        prompt = self.SUMMARIZE_PROMPT.format(text=text[:3000])
        try:
            return self._call_api([{"role": "user", "content": prompt}]).strip()
        except Exception as exc:
            logger.warning("LLM summarize failed: %s", exc)
            raise

    def extract_structured(self, text: str, schema_name: str, correlation_id: str = "") -> dict[str, Any]:
        from app.core.exceptions import LLMProviderError

        if not text.strip():
            return {}

        prompt = self.EXTRACTION_PROMPT.format(text=text[:6000])
        try:
            raw = self._call_api(
                [{"role": "user", "content": prompt}],
                correlation_id=correlation_id,
            )
            return self._extract_json_from_text(raw)
        except json.JSONDecodeError as exc:
            logger.warning(
                "LLM extract_structured JSON parse failed provider=groq model=%s error_type=%s error_message=%s corr=%s raw_len=%d",
                self.model, type(exc).__name__, str(exc)[:200], correlation_id or "-", len(raw) if "raw" in locals() else 0,
            )
            raise LLMProviderError(
                f"LLM returned invalid JSON for extraction: {exc}",
                provider="openrouter",
                model=self.model,
            ) from exc

    def classify_document(self, text: str, filename: str, declared_type: str) -> dict[str, Any]:
        from app.core.exceptions import LLMProviderError

        prompt = self.CLASSIFICATION_PROMPT.format(
            filename=filename,
            declared_type=declared_type or "UNKNOWN",
            text_excerpt=text[:1000],
        )
        try:
            raw = self._call_api([{"role": "user", "content": prompt}])
            result = self._extract_json_from_text(raw)
            return {
                "document_type": result.get("document_type", "OTHER"),
                "confidence": float(result.get("confidence", 0.5)),
                "reason": result.get("reason", ""),
                "signals": result.get("signals", []),
                "provider": "openrouter",
                "model": self.model,
            }
        except json.JSONDecodeError as exc:
            logger.warning("LLM classify_document JSON parse failed: %s | raw=%.200s", exc, raw)
            raise LLMProviderError(
                f"LLM returned invalid JSON for classification: {exc}",
                provider="openrouter",
                model=self.model,
            ) from exc


# ---------------------------------------------------------------------------
# Real Embedding Provider — Sentence Transformers
# ---------------------------------------------------------------------------


class FallbackLLMProvider(LLMProvider):
    """Try primary first, then the configured fallback model after primary exhausts retries."""

    def __init__(self, primary: GroqLLMProvider, fallback: GroqLLMProvider | None = None) -> None:
        self.primary = primary
        self.fallback = fallback
        self.model = primary.model

    def _call_with_fallback(self, operation: str, *args: Any, **kwargs: Any) -> Any:
        from app.core.exceptions import LLMProviderError

        try:
            return getattr(self.primary, operation)(*args, **kwargs)
        except LLMProviderError as primary_exc:
            if self.fallback is None:
                raise
            logger.warning(
                "[LLM] primary_failed model=%s fallback=%s operation=%s status_code=%s",
                self.primary.model,
                self.fallback.model,
                operation,
                primary_exc.status_code or "unknown",
            )
            try:
                return getattr(self.fallback, operation)(*args, **kwargs)
            except LLMProviderError as fallback_exc:
                logger.warning(
                    "[LLM] model=%s attempt=final status=FAILED operation=%s status_code=%s",
                    self.fallback.model,
                    operation,
                    fallback_exc.status_code or "unknown",
                )
                raise LLMProviderError(
                    (
                        f"Primary model ({self.primary.model}) and fallback model "
                        f"({self.fallback.model}) failed. Primary: {primary_exc.message}. "
                        f"Fallback: {fallback_exc.message}"
                    ),
                    provider=fallback_exc.provider or primary_exc.provider or "openrouter",
                    status_code=fallback_exc.status_code,
                    retry_after=fallback_exc.retry_after,
                    model=self.fallback.model,
                ) from fallback_exc

    def generate(self, prompt: str, system: str = "") -> str:
        return self._call_with_fallback("generate", prompt, system)

    def summarize(self, text: str) -> str:
        return self._call_with_fallback("summarize", text)

    def extract_structured(self, text: str, schema_name: str, **kwargs: Any) -> dict[str, Any]:
        return self._call_with_fallback("extract_structured", text, schema_name, **kwargs)

    def classify_document(self, text: str, filename: str, declared_type: str) -> dict[str, Any]:
        return self._call_with_fallback("classify_document", text, filename, declared_type)


class SentenceTransformerEmbeddingProvider(EmbeddingProvider):
    """
    Real embedding using sentence-transformers.

    Falls back to LocalFallbackEmbeddingProvider with a logged warning
    only when explicitly configured for development.
    Never silently falls back in production.
    """

    def __init__(self, model_name: str = "all-MiniLM-L6-v2") -> None:
        self.model_name = model_name
        self._model: Any = None

    def _get_model(self) -> Any:
        if self._model is not None:
            return self._model
        try:
            from sentence_transformers import SentenceTransformer  # type: ignore
            self._model = SentenceTransformer(self.model_name)
            return self._model
        except ImportError as exc:
            from app.core.exceptions import EmbeddingProviderError
            raise EmbeddingProviderError(
                "sentence-transformers is not installed. "
                "Install with: pip install sentence-transformers",
                provider="sentence_transformers",
            ) from exc

    def embed(self, texts: list[str]) -> list[list[float]]:
        model = self._get_model()
        embeddings = model.encode(texts, convert_to_numpy=True)
        return [emb.tolist() for emb in embeddings]


class LocalFallbackEmbeddingProvider(EmbeddingProvider):
    """
    Deterministic hash-based embedding for LOCAL/DEV mode ONLY.

    NOT suitable for semantic retrieval. 
    Used only when EMBEDDING_PROVIDER=local and environment is development.
    """

    def embed(self, texts: list[str]) -> list[list[float]]:
        vectors: list[list[float]] = []
        for text in texts:
            buckets = [0.0] * 64  # Larger vector for slightly better separation
            for index, char in enumerate(text.lower()):
                buckets[index % 64] += (ord(char) % 37) / 37.0
            total = sum(buckets) or 1.0
            vectors.append([round(v / total, 6) for v in buckets])
        return vectors


# ---------------------------------------------------------------------------
# Provider Factory
# ---------------------------------------------------------------------------


def get_llm_provider(settings: Any | None = None) -> LLMProvider:
    """
    Return the configured production LLM provider.
    Supports: openai_compatible, openrouter, groq (legacy alias).
    Raises ConfigurationError if credentials are missing in non-demo mode.
    """
    from app.core.config import get_settings
    from app.core.exceptions import ConfigurationError

    s = settings or get_settings()

    if s.llm_provider in ("groq", "openai_compatible", "openrouter"):
        if not s.llm_api_key:
            if s.demo_mode:
                return _UnconfiguredLLMProvider(s.llm_provider)
            raise ConfigurationError(
                f"OPENROUTER_API_KEY is required when LLM_PROVIDER='{s.llm_provider}' "
                "and DEMO_MODE=false. Set OPENROUTER_API_KEY in your .env file."
            )
        primary = GroqLLMProvider(
            api_key=s.llm_api_key,
            model=s.llm_model,
            base_url=s.llm_base_url,
            temperature=s.llm_temperature,
            timeout=s.llm_timeout,
            max_tokens=s.llm_max_tokens,
            max_retries=s.llm_max_retries,
            retry_backoff_seconds=s.llm_retry_backoff_seconds,
        )
        fallback = None
        if s.llm_fallback_model and s.llm_fallback_model != s.llm_model:
            fallback = GroqLLMProvider(
                api_key=s.llm_api_key,
                model=s.llm_fallback_model,
                base_url=s.llm_base_url,
                temperature=s.llm_temperature,
                timeout=s.llm_timeout,
                max_tokens=s.llm_max_tokens,
                max_retries=s.llm_max_retries,
                retry_backoff_seconds=s.llm_retry_backoff_seconds,
            )
        return FallbackLLMProvider(primary, fallback) if fallback else primary

    raise ConfigurationError(f"Unsupported LLM_PROVIDER: '{s.llm_provider}'. Supported: openrouter, openai_compatible")


def get_ocr_provider(settings: Any | None = None) -> OCRProvider | None:
    """Return the configured OCR provider, or None if OCR is disabled."""
    from app.core.config import get_settings

    s = settings or get_settings()
    if not s.ocr_enabled or s.ocr_provider == "none":
        return None
    if s.ocr_provider == "tesseract":
        return TesseractOCRProvider(
            tesseract_cmd=s.tesseract_cmd,
            language=s.ocr_language,
            timeout=s.ocr_timeout,
        )
    from app.core.exceptions import ConfigurationError
    raise ConfigurationError(f"Unsupported OCR_PROVIDER: '{s.ocr_provider}'. Supported: tesseract, none")


def get_embedding_provider(settings: Any | None = None) -> EmbeddingProvider:
    """Return the configured embedding provider."""
    from app.core.config import get_settings

    s = settings or get_settings()
    if s.embedding_provider == "sentence_transformers":
        return SentenceTransformerEmbeddingProvider(model_name=s.embedding_model)
    if s.embedding_provider == "local":
        logger.warning(
            "EMBEDDING_PROVIDER=local uses deterministic hash embeddings. "
            "NOT suitable for production semantic search. "
            "Set EMBEDDING_PROVIDER=sentence_transformers for real semantic retrieval."
        )
        return LocalFallbackEmbeddingProvider()
    from app.core.exceptions import ConfigurationError
    raise ConfigurationError(f"Unsupported EMBEDDING_PROVIDER: '{s.embedding_provider}'")


class _UnconfiguredLLMProvider(LLMProvider):
    """Placeholder used when demo_mode=True but API key is missing. Raises on all calls."""

    def __init__(self, provider: str) -> None:
        self.provider = provider

    def _raise(self) -> None:
        from app.core.exceptions import LLMProviderError
        raise LLMProviderError(
            f"LLM provider '{self.provider}' is not configured (no API key). "
            "Set OPENROUTER_API_KEY in your .env file.",
            provider=self.provider,
        )

    def generate(self, prompt: str, system: str = "") -> str:
        self._raise()
        return ""  # unreachable

    def summarize(self, text: str) -> str:
        self._raise()
        return ""

    def extract_structured(self, text: str, schema_name: str, **kwargs: Any) -> dict[str, Any]:
        self._raise()
        return {}

    def classify_document(self, text: str, filename: str, declared_type: str) -> dict[str, Any]:
        self._raise()
        return {}


# ---------------------------------------------------------------------------
# TEST-ONLY — DO NOT USE IN PRODUCTION CODE
# ---------------------------------------------------------------------------

# Legacy alias — kept so existing test imports don't break
